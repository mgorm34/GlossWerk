"""
GlossWerk - Document Processing Pipeline
Accepts a .docx patent document in German, runs it through:
  1. DeepL translation
  2. T5 APE correction
  3. CometKiwi QE scoring
  4. Segment triage (publish / review / full_edit)
  5. Outputs annotated .docx with color-coded segments

Usage:
    python 12_process_document.py --input patent.docx --model "C:\glosswerk\models\patent_ape_stageA\final"
    python 12_process_document.py --input patent.docx --model "..." --glossary my_terms.tsv
    python 12_process_document.py --input patent.docx --model "..." --deepl_key YOUR_KEY
"""

import argparse
import csv
import json
import os
import re
import sys
import time
from datetime import datetime

import torch
from transformers import T5ForConditionalGeneration, T5TokenizerFast

# Optional imports - graceful handling
try:
    import deepl
    HAS_DEEPL = True
except ImportError:
    HAS_DEEPL = False

try:
    from comet import download_model, load_from_checkpoint
    HAS_COMET = True
except ImportError:
    HAS_COMET = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(PROJECT_ROOT, "data")


def check_dependencies():
    """Check all required packages are available."""
    missing = []
    if not HAS_DEEPL:
        missing.append("deepl (pip install deepl)")
    if not HAS_COMET:
        missing.append("unbabel-comet (pip install unbabel-comet)")
    if not HAS_DOCX:
        missing.append("python-docx (pip install python-docx)")
    if missing:
        print("Missing dependencies:")
        for m in missing:
            print(f"  pip install {m}")
        sys.exit(1)


def extract_sentences_from_docx(filepath):
    """Extract text from a .docx file, split into sentences."""
    doc = DocxDocument(filepath)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Join all paragraphs and split into sentences
    combined = " ".join(full_text)

    # Simple sentence splitting - handles German/English patent text
    # Split on period followed by space and uppercase, or on newlines
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ])', combined)

    # Filter out very short fragments
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

    return sentences


def translate_with_deepl(sentences, api_key, source_lang="DE", target_lang="EN-US"):
    """Translate German sentences using DeepL API."""
    translator = deepl.Translator(api_key)

    print(f"  Translating {len(sentences)} sentences with DeepL...")
    translations = []
    batch_size = 50

    for i in range(0, len(sentences), batch_size):
        batch = sentences[i:i + batch_size]
        try:
            results = translator.translate_text(batch, source_lang=source_lang, target_lang=target_lang)
            for r in results:
                translations.append(r.text)
        except Exception as e:
            print(f"  DeepL error at batch {i}: {e}")
            # Fall back to empty translations for failed batch
            translations.extend(["[TRANSLATION FAILED]"] * len(batch))
        time.sleep(0.3)  # Rate limiting

        if (i + batch_size) % 200 == 0:
            print(f"  Translated {min(i + batch_size, len(sentences))}/{len(sentences)}")

    return translations


def load_glossary(filepath):
    """Load terminology glossary from TSV file."""
    terms = {}
    if not filepath or not os.path.exists(filepath):
        return terms

    with open(filepath, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f, delimiter="\t")
        for row in reader:
            de = row.get("de_term", row.get("German", "")).strip()
            en = row.get("correct_en", row.get("Correct EN", row.get("en_term", ""))).strip()
            if de and en:
                terms[de.lower()] = (de, en)

    print(f"  Loaded {len(terms)} glossary terms")
    return terms


def match_glossary_terms(source_sentence, glossary, max_terms=5):
    """Find glossary terms in a German source sentence."""
    if not glossary:
        return []

    source_lower = source_sentence.lower()
    matches = []

    # Sort by term length (longest first) for greedy matching
    for key in sorted(glossary.keys(), key=len, reverse=True):
        if key in source_lower and len(matches) < max_terms:
            de_term, en_term = glossary[key]
            matches.append(f"{de_term}={en_term}")

    return matches


def format_ape_input(translation, term_matches=None):
    """Format input for the APE model."""
    input_text = f"postedit: {translation}"
    if term_matches:
        input_text += " || terms: " + "; ".join(term_matches)
    return input_text


def run_ape(model, tokenizer, translations, source_sentences=None, glossary=None,
            device="cuda", batch_size=16, max_length=256):
    """Run APE corrections on translations."""
    corrections = []

    for i in range(0, len(translations), batch_size):
        batch_translations = translations[i:i + batch_size]

        # Build inputs with optional terminology
        inputs = []
        for j, trans in enumerate(batch_translations):
            term_matches = []
            if glossary and source_sentences:
                src_idx = i + j
                if src_idx < len(source_sentences):
                    term_matches = match_glossary_terms(source_sentences[src_idx], glossary)

            inputs.append(format_ape_input(trans, term_matches if term_matches else None))

        tokenized = tokenizer(
            inputs, max_length=max_length, padding=True,
            truncation=True, return_tensors="pt",
        ).to(device)

        with torch.no_grad():
            outputs = model.generate(
                **tokenized, max_length=max_length, num_beams=4, early_stopping=True,
            )

        decoded = tokenizer.batch_decode(outputs, skip_special_tokens=True)
        corrections.extend(decoded)

        if (i + batch_size) % 100 == 0 and i > 0:
            print(f"  APE: {min(i + batch_size, len(translations))}/{len(translations)}")

    return corrections


def run_qe(qe_model, sources, translations, batch_size=32):
    """Score translations with CometKiwi."""
    data = [{"src": src, "mt": mt} for src, mt in zip(sources, translations)]
    output = qe_model.predict(data, batch_size=batch_size, gpus=1 if torch.cuda.is_available() else 0)
    return output.scores


def triage(scores, threshold_high=0.82, threshold_low=0.68):
    """Assign confidence buckets."""
    labels = []
    for s in scores:
        if s >= threshold_high:
            labels.append("publish")
        elif s >= threshold_low:
            labels.append("review")
        else:
            labels.append("full_edit")
    return labels


def create_output_docx(segments, output_path):
    """Create annotated output document with color-coded segments."""
    doc = DocxDocument()

    # Title
    title = doc.add_paragraph()
    run = title.add_run("GlossWerk - Patent Translation Analysis")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

    # Summary
    summary = doc.add_paragraph()
    total = len(segments)
    publish_count = sum(1 for s in segments if s["triage"] == "publish")
    review_count = sum(1 for s in segments if s["triage"] == "review")
    edit_count = sum(1 for s in segments if s["triage"] == "full_edit")
    changed_count = sum(1 for s in segments if s["changed"])

    run = summary.add_run(
        f"Processed {total} segments  |  "
        f"Publish: {publish_count} ({100*publish_count/total:.0f}%)  |  "
        f"Review: {review_count} ({100*review_count/total:.0f}%)  |  "
        f"Full Edit: {edit_count} ({100*edit_count/total:.0f}%)  |  "
        f"APE Changed: {changed_count} ({100*changed_count/total:.0f}%)"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Timestamp
    ts = doc.add_paragraph()
    run = ts.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # Legend
    legend = doc.add_paragraph()
    run = legend.add_run("Legend: ")
    run.bold = True
    run.font.size = Pt(10)

    run = legend.add_run("\u2588 PUBLISH ")
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    run.font.size = Pt(10)

    run = legend.add_run("\u2588 REVIEW ")
    run.font.color.rgb = RGBColor(0xF5, 0x7F, 0x17)
    run.font.size = Pt(10)

    run = legend.add_run("\u2588 FULL EDIT")
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Segments
    for i, seg in enumerate(segments):
        # Segment header
        header = doc.add_paragraph()
        triage_label = seg["triage"].upper()
        qe_score = seg["qe_score"]

        if seg["triage"] == "publish":
            color = RGBColor(0x2E, 0x7D, 0x32)
        elif seg["triage"] == "review":
            color = RGBColor(0xF5, 0x7F, 0x17)
        else:
            color = RGBColor(0xC6, 0x28, 0x28)

        run = header.add_run(f"[{i+1}] {triage_label} (QE: {qe_score:.3f})")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = color

        if seg["changed"]:
            run = header.add_run("  \u2022 APE corrected")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        if seg.get("terms_applied"):
            run = header.add_run(f"  \u2022 Terms: {', '.join(seg['terms_applied'])}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6A, 0x1B, 0x9A)

        # Source
        src_para = doc.add_paragraph()
        run = src_para.add_run("DE: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run = src_para.add_run(seg["source"])
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Translation (the corrected output)
        trans_para = doc.add_paragraph()
        run = trans_para.add_run("EN: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = color
        run = trans_para.add_run(seg["corrected"])
        run.font.size = Pt(10)

        # If changed, show original DeepL for comparison
        if seg["changed"]:
            orig_para = doc.add_paragraph()
            run = orig_para.add_run("DeepL original: ")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run = orig_para.add_run(seg["deepl"])
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # Spacer between segments
        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def create_json_output(segments, output_path):
    """Create JSON output for programmatic consumption."""
    output = {
        "generated": datetime.now().isoformat(),
        "total_segments": len(segments),
        "summary": {
            "publish": sum(1 for s in segments if s["triage"] == "publish"),
            "review": sum(1 for s in segments if s["triage"] == "review"),
            "full_edit": sum(1 for s in segments if s["triage"] == "full_edit"),
            "ape_changed": sum(1 for s in segments if s["changed"]),
        },
        "segments": segments,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="GlossWerk Document Processing Pipeline")
    parser.add_argument("--input", type=str, required=True, help="Input .docx file (German)")
    parser.add_argument("--model", type=str, required=True, help="Path to APE model")
    parser.add_argument("--deepl_key", type=str, default=None,
                        help="DeepL API key (or set DEEPL_AUTH_KEY env var)")
    parser.add_argument("--glossary", type=str, default=None,
                        help="Path to terminology glossary TSV")
    parser.add_argument("--output_dir", type=str, default=None,
                        help="Output directory (default: same as input)")
    parser.add_argument("--batch_size", type=int, default=16, help="APE batch size")
    parser.add_argument("--qe_batch_size", type=int, default=32, help="QE batch size")
    parser.add_argument("--threshold_high", type=float, default=0.82)
    parser.add_argument("--threshold_low", type=float, default=0.68)
    parser.add_argument("--skip_qe", action="store_true", help="Skip QE scoring")
    args = parser.parse_args()

    check_dependencies()

    # Get API key
    deepl_key = args.deepl_key or os.environ.get("DEEPL_AUTH_KEY")
    if not deepl_key:
        print("ERROR: No DeepL API key. Use --deepl_key or set DEEPL_AUTH_KEY")
        sys.exit(1)

    # Output paths
    if args.output_dir:
        os.makedirs(args.output_dir, exist_ok=True)
        out_dir = args.output_dir
    else:
        out_dir = os.path.dirname(os.path.abspath(args.input))

    base_name = os.path.splitext(os.path.basename(args.input))[0]
    output_docx = os.path.join(out_dir, f"{base_name}_glosswerk.docx")
    output_json = os.path.join(out_dir, f"{base_name}_glosswerk.json")

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    print("=" * 60)
    print("GlossWerk - Document Processing Pipeline")
    print("=" * 60)
    print(f"Input:    {args.input}")
    print(f"Model:    {args.model}")
    print(f"Glossary: {args.glossary or 'None'}")
    print(f"Device:   {device}")
    print()

    # Step 1: Extract sentences
    print("--- Step 1: Extract sentences ---")
    sentences = extract_sentences_from_docx(args.input)
    print(f"  Extracted {len(sentences)} sentences")

    if not sentences:
        print("ERROR: No sentences found in document.")
        sys.exit(1)

    # Step 2: Translate with DeepL
    print("\n--- Step 2: DeepL Translation ---")
    translations = translate_with_deepl(sentences, deepl_key)
    print(f"  Translated {len(translations)} sentences")

    # Step 3: Load glossary
    glossary = load_glossary(args.glossary) if args.glossary else {}

    # Step 4: APE correction
    print("\n--- Step 3: APE Correction ---")
    print(f"  Loading model: {args.model}")
    tokenizer = T5TokenizerFast.from_pretrained(args.model)
    ape_model = T5ForConditionalGeneration.from_pretrained(args.model).to(device)
    ape_model.eval()

    corrections = run_ape(
        ape_model, tokenizer, translations,
        source_sentences=sentences, glossary=glossary,
        device=device, batch_size=args.batch_size,
    )
    print(f"  Corrected {len(corrections)} sentences")

    # Free GPU memory
    del ape_model
    torch.cuda.empty_cache()

    # Step 5: QE scoring
    if not args.skip_qe and HAS_COMET:
        print("\n--- Step 4: Quality Estimation ---")
        qe_model = load_cometkiwi()
        qe_scores = run_qe(qe_model, sentences, corrections, args.qe_batch_size)
        triage_labels = triage(qe_scores, args.threshold_high, args.threshold_low)
        del qe_model
        torch.cuda.empty_cache()
    else:
        print("\n--- Step 4: Quality Estimation (skipped) ---")
        qe_scores = [0.75] * len(corrections)  # Default middle score
        triage_labels = ["review"] * len(corrections)

    # Build segments
    segments = []
    for i in range(len(sentences)):
        changed = translations[i].strip() != corrections[i].strip()
        term_matches = match_glossary_terms(sentences[i], glossary) if glossary else []

        segments.append({
            "id": i + 1,
            "source": sentences[i],
            "deepl": translations[i],
            "corrected": corrections[i],
            "changed": changed,
            "qe_score": round(qe_scores[i], 4),
            "triage": triage_labels[i],
            "terms_applied": term_matches,
        })

    # Summary
    print(f"\n{'=' * 60}")
    print("PROCESSING SUMMARY")
    print(f"{'=' * 60}")
    total = len(segments)
    publish = sum(1 for s in segments if s["triage"] == "publish")
    review = sum(1 for s in segments if s["triage"] == "review")
    full_edit = sum(1 for s in segments if s["triage"] == "full_edit")
    changed = sum(1 for s in segments if s["changed"])

    print(f"  Total segments:  {total}")
    print(f"  Publish:         {publish} ({100*publish/total:.0f}%)")
    print(f"  Review:          {review} ({100*review/total:.0f}%)")
    print(f"  Full edit:       {full_edit} ({100*full_edit/total:.0f}%)")
    print(f"  APE changed:     {changed} ({100*changed/total:.0f}%)")

    if glossary:
        terms_used = sum(1 for s in segments if s["terms_applied"])
        print(f"  Terms applied:   {terms_used} segments")

    # Step 6: Generate outputs
    print(f"\n--- Step 5: Generate Output ---")

    create_output_docx(segments, output_docx)
    print(f"  Document: {output_docx}")

    create_json_output(segments, output_json)
    print(f"  JSON:     {output_json}")

    print(f"\n{'=' * 60}")
    print(f"Done! Open {output_docx} to review the triaged translation.")
    print(f"{'=' * 60}")


def load_cometkiwi():
    """Load CometKiwi model."""
    print("  Loading CometKiwi...")
    model_path = download_model("Unbabel/wmt22-cometkiwi-da")
    model = load_from_checkpoint(model_path)
    return model


if __name__ == "__main__":
    main()
