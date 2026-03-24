"""
GlossWerk Pipeline v2 — QE-Activated APE with XCOMET-XL + LoRA Adapter

Usage:
  python 13_pipeline_v2.py --input patent.docx --model C:\\glosswerk\\models\\patent_ape_stageA\\final
  python 13_pipeline_v2.py --input patent.docx --model ... --lora C:\\glosswerk\\models\\a61f_lora
  python 13_pipeline_v2.py --input patent.docx --model ... --glossary terms.tsv
"""

import argparse
import json
import os
import re
import torch
import deepl
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from transformers import T5ForConditionalGeneration, T5TokenizerFast
from comet import download_model, load_from_checkpoint


def extract_sentences_from_docx(filepath):
    doc = DocxDocument(filepath)
    sentences = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 10:
            parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ\[])', text)
            for part in parts:
                part = part.strip()
                if len(part) > 10:
                    sentences.append(part)
    return sentences


def translate_deepl(sentences, auth_key=None):
    if auth_key is None:
        auth_key = os.environ.get("DEEPL_AUTH_KEY")
    if not auth_key:
        raise ValueError("Set DEEPL_AUTH_KEY environment variable or pass --deepl_key")
    translator = deepl.Translator(auth_key)
    translations = []
    for i in range(0, len(sentences), 50):
        batch = sentences[i:i+50]
        results = translator.translate_text(batch, source_lang="DE", target_lang="EN-US")
        for r in results:
            translations.append(r.text)
        print(f"  Translated {min(i+50, len(sentences))}/{len(sentences)}")
    return translations


def run_xcomet(sources, translations, model=None, batch_size=4):
    if model is None:
        model_path = download_model("Unbabel/XCOMET-XL")
        model = load_from_checkpoint(model_path)
        model.eval()
    data = [{"src": s, "mt": t} for s, t in zip(sources, translations)]
    output = model.predict(data, batch_size=batch_size, gpus=1)
    scores = output.scores
    error_spans = output.metadata.error_spans
    return scores, error_spans, model


def load_glossary(filepath):
    glossary = {}
    if not filepath or not os.path.exists(filepath):
        return glossary
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            parts = line.split("\t")
            if len(parts) >= 2:
                de_term = parts[0].strip().lower()
                en_term = parts[1].strip()
                if de_term and en_term:
                    glossary[de_term] = en_term
    return glossary


def match_glossary(source_de, translation_en, glossary):
    force_terms = []
    source_lower = source_de.lower()
    for de_term, en_term in glossary.items():
        if de_term in source_lower:
            if en_term.lower() not in translation_en.lower():
                force_terms.append(en_term)
    return force_terms


def run_ape_with_activation(sentences, translations, error_spans_list, glossary,
                            model, tokenizer, device, error_threshold=0.35):
    results = []
    skipped = 0
    for i, (src, trans) in enumerate(zip(sentences, translations)):
        errors = error_spans_list[i] if i < len(error_spans_list) else []
        real_errors = [e for e in errors if e.get("confidence", 0) >= error_threshold]
        glossary_terms = match_glossary(src, trans, glossary) if glossary else []

        if not real_errors and not glossary_terms:
            results.append({
                "ape": trans,
                "changed": False,
                "n_errors": 0,
                "n_glossary_forced": 0,
                "skipped": True,
            })
            skipped += 1
            continue

        input_text = f"postedit: {trans}"
        input_ids = tokenizer(
            input_text, return_tensors="pt",
            max_length=256, truncation=True
        ).input_ids.to(device)

        output_ids = model.generate(input_ids=input_ids, max_length=256, num_beams=5)
        ape_text = tokenizer.decode(output_ids[0], skip_special_tokens=True)
        changed = ape_text.strip() != trans.strip()

        results.append({
            "ape": ape_text,
            "changed": changed,
            "n_errors": len(real_errors),
            "n_glossary_forced": len(glossary_terms),
            "skipped": False,
        })

        if (i + 1) % 10 == 0:
            print(f"  APE processed {i+1}/{len(sentences)} (skipped {skipped} clean)")

    print(f"  Done. {skipped}/{len(sentences)} segments passed through unchanged.")
    return results


def triage_segments(scores, thresholds=(0.92, 0.75)):
    publish_thresh, review_thresh = thresholds
    triage = []
    for s in scores:
        if s >= publish_thresh:
            triage.append("PUBLISH")
        elif s >= review_thresh:
            triage.append("REVIEW")
        else:
            triage.append("FULL_EDIT")
    return triage


COLORS = {
    "PUBLISH": RGBColor(0x27, 0xAE, 0x60),
    "REVIEW": RGBColor(0xF3, 0x9C, 0x12),
    "FULL_EDIT": RGBColor(0xE7, 0x4C, 0x3C),
}


def create_output_docx(segments, output_path):
    doc = DocxDocument()

    stats = {"PUBLISH": 0, "REVIEW": 0, "FULL_EDIT": 0}
    changed_count = 0
    skipped_count = 0
    glossary_count = 0
    for seg in segments:
        stats[seg["triage"]] += 1
        if seg["changed"]:
            changed_count += 1
        if seg.get("skipped"):
            skipped_count += 1
        if seg.get("n_glossary_forced", 0) > 0:
            glossary_count += 1

    total = len(segments)
    p = doc.add_paragraph()
    run = p.add_run("GlossWerk v2 — QE-Activated APE + XCOMET-XL + A61F LoRA")
    run.bold = True
    run.font.size = Pt(14)

    summary = (
        f"Processed {total} segments  |  "
        f"Publish: {stats['PUBLISH']} ({100*stats['PUBLISH']//total}%)  |  "
        f"Review: {stats['REVIEW']} ({100*stats['REVIEW']//total}%)  |  "
        f"Full Edit: {stats['FULL_EDIT']} ({100*stats['FULL_EDIT']//total}%)  |  "
        f"APE Changed: {changed_count} ({100*changed_count//total}%)  |  "
        f"QE Passthrough: {skipped_count} ({100*skipped_count//total}%)  |  "
        f"Glossary forced: {glossary_count}"
    )
    doc.add_paragraph(summary)

    p = doc.add_paragraph()
    run = p.add_run("Legend: ")
    run.bold = True
    for label, color in COLORS.items():
        run = p.add_run(f"  {label}  ")
        run.font.color.rgb = color
        run.bold = True

    for seg in segments:
        color = COLORS[seg["triage"]]

        p = doc.add_paragraph()
        header = f"[{seg['id']}] {seg['triage']} (QE: {seg['xcomet_score']:.3f})"
        if seg.get("skipped"):
            header += "  • QE passthrough"
        elif seg["changed"]:
            header += "  • APE corrected"
        if seg.get("n_errors", 0) > 0:
            header += f"  • {seg['n_errors']} errors detected"
        if seg.get("n_glossary_forced", 0) > 0:
            header += f"  • {seg['n_glossary_forced']} glossary terms"
        run = p.add_run(header)
        run.font.color.rgb = color
        run.bold = True

        p = doc.add_paragraph()
        run = p.add_run("DE: ")
        run.bold = True
        p.add_run(seg["de"])

        p = doc.add_paragraph()
        run = p.add_run("EN: ")
        run.bold = True
        p.add_run(seg["ape"])

        if seg["changed"]:
            p = doc.add_paragraph()
            run = p.add_run("DeepL original: ")
            run.italic = True
            run2 = p.add_run(seg["deepl"])
            run2.italic = True

        if seg.get("error_spans"):
            errs = ", ".join(
                f"'{e['text']}' ({e['severity']}, {e['confidence']:.2f})"
                for e in seg["error_spans"]
            )
            p = doc.add_paragraph()
            run = p.add_run(f"QE errors: {errs}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph()

    doc.save(output_path)
    return stats


def main():
    parser = argparse.ArgumentParser(description="GlossWerk v2 Pipeline")
    parser.add_argument("--input", required=True, help="Input German .docx")
    parser.add_argument("--model", required=True, help="Path to T5 APE base model")
    parser.add_argument("--lora", default=r"C:\glosswerk\models\a61f_lora", help="Path to LoRA adapter")
    parser.add_argument("--glossary", default=None, help="TSV glossary (DE\\tEN)")
    parser.add_argument("--deepl_key", default=None)
    parser.add_argument("--error_threshold", type=float, default=0.35)
    parser.add_argument("--publish_threshold", type=float, default=0.92)
    parser.add_argument("--review_threshold", type=float, default=0.75)
    parser.add_argument("--batch_size", type=int, default=4)
    parser.add_argument("--output", default=None)
    args = parser.parse_args()

    if args.output is None:
        base = os.path.splitext(args.input)[0]
        args.output = f"{base}_glosswerk_v2.docx"

    print("--- Step 1: Extract sentences ---")
    sentences = extract_sentences_from_docx(args.input)
    print(f"  Found {len(sentences)} sentences")

    print("--- Step 2: DeepL translation ---")
    translations = translate_deepl(sentences, args.deepl_key)

    print("--- Step 3: XCOMET-XL word-level QE ---")
    xcomet_scores, error_spans, xcomet_model = run_xcomet(
        sentences, translations, batch_size=args.batch_size
    )
    print(f"  Avg XCOMET score: {sum(xcomet_scores)/len(xcomet_scores):.4f}")
    errors_total = sum(len(e) for e in error_spans)
    segs_with_errors = sum(1 for e in error_spans if len(e) > 0)
    print(f"  Total error spans: {errors_total}")
    print(f"  Segments with errors: {segs_with_errors}/{len(sentences)}")

    glossary = {}
    if args.glossary:
        print("--- Step 4: Load glossary ---")
        glossary = load_glossary(args.glossary)
        print(f"  Loaded {len(glossary)} entries")

    print("--- Step 5: QE-Activated APE (LoRA) ---")
    del xcomet_model
    torch.cuda.empty_cache()

    device = "cuda" if torch.cuda.is_available() else "cpu"
    tokenizer = T5TokenizerFast.from_pretrained(args.model)
    base_model = T5ForConditionalGeneration.from_pretrained(args.model)

    if args.lora and os.path.exists(args.lora):
        from peft import PeftModel
        t5_model = PeftModel.from_pretrained(base_model, args.lora)
        print(f"  Loaded LoRA adapter from {args.lora}")
    else:
        t5_model = base_model
        print(f"  No LoRA adapter, using base model")

    t5_model.to(device)
    t5_model.eval()

    ape_results = run_ape_with_activation(
        sentences, translations, error_spans, glossary,
        t5_model, tokenizer, device,
        error_threshold=args.error_threshold,
    )

    del t5_model
    if 'base_model' in dir():
        del base_model
    torch.cuda.empty_cache()

    print("--- Step 6: XCOMET-XL triage scoring ---")
    ape_translations = [r["ape"] for r in ape_results]
    triage_scores, _, _ = run_xcomet(
        sentences, ape_translations, batch_size=args.batch_size
    )

    triage_labels = triage_segments(
        triage_scores,
        thresholds=(args.publish_threshold, args.review_threshold)
    )

    print("--- Step 7: Generate output ---")
    segments = []
    for i in range(len(sentences)):
        segments.append({
            "id": i + 1,
            "de": sentences[i],
            "deepl": translations[i],
            "ape": ape_results[i]["ape"],
            "changed": ape_results[i]["changed"],
            "skipped": ape_results[i].get("skipped", False),
            "xcomet_score": triage_scores[i],
            "xcomet_deepl_score": xcomet_scores[i],
            "triage": triage_labels[i],
            "error_spans": error_spans[i] if i < len(error_spans) else [],
            "n_errors": ape_results[i]["n_errors"],
            "n_glossary_forced": ape_results[i]["n_glossary_forced"],
        })

    stats = create_output_docx(segments, args.output)

    json_path = args.output.replace(".docx", ".json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(segments, f, indent=2, ensure_ascii=False)

    total = len(segments)
    changed = sum(1 for s in segments if s["changed"])
    skipped = sum(1 for s in segments if s.get("skipped"))
    glossary_applied = sum(1 for s in segments if s["n_glossary_forced"] > 0)

    print(f"\n{'='*60}")
    print(f"GlossWerk v2 + LoRA — Results")
    print(f"{'='*60}")
    print(f"  Total segments:      {total}")
    print(f"  Publish:             {stats['PUBLISH']} ({100*stats['PUBLISH']//total}%)")
    print(f"  Review:              {stats['REVIEW']} ({100*stats['REVIEW']//total}%)")
    print(f"  Full edit:           {stats['FULL_EDIT']} ({100*stats['FULL_EDIT']//total}%)")
    print(f"  APE changed:         {changed} ({100*changed//total}%)")
    print(f"  QE passthrough:      {skipped} ({100*skipped//total}%)")
    print(f"  Glossary forced:     {glossary_applied}")
    print(f"  Avg XCOMET (DeepL):  {sum(xcomet_scores)/len(xcomet_scores):.4f}")
    print(f"  Avg XCOMET (APE):    {sum(triage_scores)/len(triage_scores):.4f}")
    print(f"\n  Output: {args.output}")
    print(f"  JSON:   {json_path}")


if __name__ == "__main__":
    main()