"""
GlossWerk — DE→EN Translation Pipeline for Technical Text & Patents

Upload a German patent → scan & review terminology → translate with QE →
review & edit side-by-side → export.

Run: streamlit run glosswerk_app.py
Requires: ANTHROPIC_API_KEY environment variable
Install: pip install anthropic streamlit python-docx spacy
         python -m spacy download de_core_news_lg
"""

import json
import os
import re
import sys
import tempfile
import time
from collections import Counter
from pathlib import Path

import streamlit as st
from docx.shared import Pt

# --- Path setup ---
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# --- Module imports ---
import importlib.util

_spec = importlib.util.spec_from_file_location(
    "extract_terms",
    os.path.join(PROJECT_ROOT, "skills", "glosswerk-term-scanner", "scripts", "extract_terms.py"),
)
_mod = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_mod)

extract_text_from_docx_terms = _mod.extract_text_from_docx
extract_sentences = _mod.extract_sentences
extract_nouns_spacy = _mod.extract_nouns_spacy
extract_nouns_heuristic = _mod.extract_nouns_heuristic
extract_technical_adjectives = _mod.extract_technical_adjectives
extract_patent_verbs = _mod.extract_patent_verbs
HAS_SPACY = _mod.HAS_SPACY

_spec2 = importlib.util.spec_from_file_location(
    "analyze_structure",
    os.path.join(PROJECT_ROOT, "skills", "glosswerk-structural-analyzer", "scripts", "analyze_structure.py"),
)
_mod2 = importlib.util.module_from_spec(_spec2)
_spec2.loader.exec_module(_mod2)
analyze_document = _mod2.analyze_document

from translate import (
    translate_document, extract_text_from_docx as extract_text_translate,
    split_sentences as split_sentences_translate,
)
from quality_estimate import (
    evaluate_translations, compute_triage,
    load_training_pairs, select_few_shot_examples,
)
from prompt_layers import get_available_domains
from assemble import assemble_document
from demo_auth import show_auth_gate, record_patent_use, validate_code, WATERMARK_TEXT

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

# Demo mode toggle — set to False for local dev without auth
DEMO_MODE = os.environ.get("GLOSSWERK_DEMO", "false").lower() == "true"


@st.cache_resource
def load_nlp():
    try:
        import spacy
        return spacy.load("de_core_news_lg")
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG & STYLING
# ══════════════════════════════════════════════════════════════════════════════

# Load logo for favicon
_favicon_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "logo.svg")
if os.path.exists(_favicon_path):
    st.set_page_config(page_title="GlossWerk", page_icon=_favicon_path, layout="wide")
else:
    st.set_page_config(page_title="GlossWerk", page_icon="G", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;700&family=JetBrains+Mono:wght@400;500&display=swap');
    .stApp { font-family: 'DM Sans', sans-serif; }
    h1, h2, h3 { font-family: 'DM Sans', sans-serif; font-weight: 700; }

    .main-header {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
        padding: 1.5rem 2rem; border-radius: 12px; margin-bottom: 1.5rem; color: white;
        display: flex; align-items: center; gap: 1.2rem;
    }
    .main-header .logo-icon { width: 48px; height: 48px; border-radius: 10px; flex-shrink: 0; }
    .main-header img { display: none; }
    .main-header h1 {
        color: white; margin: 0; font-size: 2rem; letter-spacing: -0.5px;
        font-weight: 700; line-height: 1.1;
    }
    .main-header h1 .gw-g { color: #10b981; font-weight: 800; }
    .main-header h1 .gw-loss { color: #e2e8f0; font-weight: 400; }
    .main-header h1 .gw-werk { color: #10b981; font-weight: 700; }
    .main-header p { color: #94a3b8; margin: 0.3rem 0 0 0; font-size: 0.95rem; }

    .triage-card { padding: 1rem 1.2rem; border-radius: 10px; text-align: center; font-weight: 600; }
    .triage-green { background: #d1fae5; color: #065f46; border: 2px solid #6ee7b7; }
    .triage-orange { background: #fef3c7; color: #92400e; border: 2px solid #fcd34d; }
    .triage-red { background: #fee2e2; color: #991b1b; border: 2px solid #fca5a5; }

    .conf-high { color: #059669; font-weight: 600; }
    .conf-medium { color: #d97706; font-weight: 600; }
    .conf-low { color: #dc2626; font-weight: 600; }

    .qe-note { font-size: 0.88rem; padding: 0.6rem 0.8rem; border-radius: 6px; margin-top: 0.4rem; color: #1f2937; }
    .qe-suggestion { font-size: 1rem; line-height: 1.6; margin-top: 0.5rem; padding: 0.5rem 0; }
    .qe-suggestion strong { background: #fef08a; padding: 1px 3px; border-radius: 2px; }
    .qe-minor { background: #fef9c3; border-left: 3px solid #eab308; color: #713f12; }
    .qe-major { background: #fee2e2; border-left: 3px solid #ef4444; color: #7f1d1d; }
    .qe-critical { background: #fecaca; border-left: 3px solid #b91c1c; color: #7f1d1d; }
    .qe-good { background: #d1fae5; border-left: 3px solid #10b981; color: #065f46; }

    .source-box {
        font-size: 0.9rem; color: #374151; padding: 0.6rem;
        background: #f8fafc; border-radius: 6px; border: 1px solid #e2e8f0;
        min-height: 60px; line-height: 1.5;
    }
    .locked-box {
        font-size: 0.9rem; color: #065f46; padding: 0.6rem;
        background: #f0fdf4; border-radius: 6px; border: 1px solid #6ee7b7;
        min-height: 60px; line-height: 1.5;
    }

    .reasoning-text { font-size: 0.8rem; color: #6b7280; font-style: italic; margin-top: 0.2rem; }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    [data-testid="stSidebar"] { min-width: 280px; max-width: 320px; }
</style>
""", unsafe_allow_html=True)

# Load logo as base64 for inline embedding
import base64 as _b64
_logo_path = os.path.join(PROJECT_ROOT, "assets", "logo.svg")
_logo_b64 = ""
if os.path.exists(_logo_path):
    with open(_logo_path, "rb") as _lf:
        _logo_b64 = _b64.b64encode(_lf.read()).decode()

_logo_img = f'<img src="data:image/svg+xml;base64,{_logo_b64}" alt="GlossWerk">' if _logo_b64 else ""

_inline_logo_svg = '<svg class="logo-icon" xmlns="http://www.w3.org/2000/svg" viewBox="0 0 512 512" fill="none"><rect width="512" height="512" rx="96" fill="#1a1a2e"/><path d="M310 140H240c-55.228 0-100 44.772-100 100v32c0 55.228 44.772 100 100 100h50v-116h-50" stroke="#10b981" stroke-width="44" stroke-linecap="round" stroke-linejoin="round" fill="none"/><line x1="284" y1="256" x2="380" y2="256" stroke="#10b981" stroke-width="44" stroke-linecap="round"/><polyline points="348,196 414,256 348,316" stroke="#10b981" stroke-width="36" stroke-linecap="round" stroke-linejoin="round" fill="none"/></svg>'

st.markdown(f"""
<div class="main-header">
    {_inline_logo_svg}
    <div>
        <h1><span class="gw-g">G</span><span class="gw-loss">loss</span><span class="gw-werk">Werk</span></h1>
        <p>DE &rarr; EN Translation Pipeline</p>
    </div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# DEMO AUTH GATE
# ══════════════════════════════════════════════════════════════════════════════

demo_auth = None
if DEMO_MODE:
    demo_auth = show_auth_gate()
    if not demo_auth:
        st.stop()
    # Show demo status bar
    st.markdown(
        f"<div style='background:#f0fdf4; padding:0.5rem 1rem; border-radius:8px; "
        f"font-size:0.85rem; color:#065f46; border:1px solid #6ee7b7; margin-bottom:1rem;'>"
        f"Demo — <strong>{demo_auth['company']}</strong> · "
        f"{demo_auth['patents_remaining']} patents remaining · "
        f"{demo_auth['days_remaining']} days left</div>",
        unsafe_allow_html=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("### Settings")
    # In demo mode, API key is server-side — hide from user
    if DEMO_MODE:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        model = "claude-sonnet-4-6"
    else:
        api_key = st.text_input("API Key", value=os.environ.get("ANTHROPIC_API_KEY", ""),
                                type="password")
        model = st.selectbox("Model", ["claude-sonnet-4-6", "claude-opus-4-6",
                                        "claude-haiku-4-5-20251001"], index=0)

    # Hardcoded pipeline defaults (no user-facing batch controls)
    batch_size_trans = 50
    batch_size_qe = 10
    n_few_shot = 30
    training_pairs_path = os.path.join(PROJECT_ROOT, "data", "hter_training", "training_pairs.jsonl")
    min_adj_freq = 2  # default, overridden in terminology tab

    st.divider()
    st.markdown("### Domain")
    _domains = get_available_domains()
    _domain_keys = list(_domains.keys())
    _domain_labels = list(_domains.values())
    selected_domain = st.selectbox(
        "Translation domain",
        options=_domain_keys,
        format_func=lambda k: _domains[k],
        index=0,
        help="Selects domain-specific translation and QE rules. Patent is the default."
    )

    st.divider()
    st.markdown("### Document")
    uploaded_file = st.file_uploader("German text (.docx)", type=["docx"])
    glossary_upload = st.file_uploader("Glossary TSV (optional)", type=["tsv", "txt"])


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════

_defaults = {
    "raw_text": None, "sentences": None, "structural_analysis": None,
    "glossary": {}, "noun_counts": None, "adj_counts": None,
    "adj_variants": None, "verb_info": None, "lemma_map": None,
    "noun_proposals": {}, "adj_proposals": {},
    "translations": None, "qe_results": None, "triage": None,
    "confirmed": {},
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v if not isinstance(v, dict) else dict(v)

# Parse uploaded glossary
pre_glossary = {}
if glossary_upload:
    content = glossary_upload.read().decode("utf-8")
    for line in content.strip().split("\n"):
        parts = line.strip().split("\t")
        if len(parts) >= 2:
            pre_glossary[parts[0].strip()] = parts[1].strip()

for de, en in pre_glossary.items():
    if de not in st.session_state.glossary:
        st.session_state.glossary[de] = en

if uploaded_file is not None:
    # Save uploaded file bytes to session state (survives reruns even if widget resets)
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    # Only re-save if it's a different file (avoid rewriting on every rerun)
    if "docx_bytes" not in st.session_state or st.session_state.get("docx_name") != uploaded_file.name:
        st.session_state["docx_bytes"] = file_bytes
        st.session_state["docx_name"] = uploaded_file.name
        # Clear previous results when a new file is uploaded (reset to defaults)
        for key in ["raw_text", "sentences", "noun_counts", "adj_counts", "adj_variants",
                     "verb_info", "lemma_map", "translations", "qe_results", "triage",
                     "structural_analysis", "doc_sentence_count"]:
            st.session_state[key] = None
        for key in ["noun_proposals", "adj_proposals"]:
            st.session_state[key] = {}
        st.session_state["glossary"] = {}
        st.session_state["confirmed"] = {}

# Reconstruct temp file from stored bytes on every run
if "docx_bytes" not in st.session_state:
    st.info("Upload a German .docx file in the sidebar to get started.")
    st.stop()

_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
_tmp.write(st.session_state["docx_bytes"])
_tmp.close()
docx_path = _tmp.name

# ══════════════════════════════════════════════════════════════════════════════
# REFRESH WARNING — warn user before losing translation progress
# ══════════════════════════════════════════════════════════════════════════════

if st.session_state.get("translations") is not None:
    import streamlit.components.v1 as components
    components.html("""
    <script>
    window.addEventListener('beforeunload', function(e) {
        e.preventDefault();
        e.returnValue = 'You have unsaved translation progress. Are you sure you want to leave?';
        return e.returnValue;
    });
    </script>
    """, height=0)


# ══════════════════════════════════════════════════════════════════════════════
# HELPER: term translation with alternatives
# ══════════════════════════════════════════════════════════════════════════════

def translate_terms_batch(terms, term_type="noun"):
    """Get EN translations with alternatives for a list of DE terms."""
    if not terms or not api_key or not HAS_ANTHROPIC:
        return {}

    client = anthropic.Anthropic(api_key=api_key)
    all_proposals = {}

    if term_type == "noun":
        schema = (
            '{"de": "German", "en": "best translation", '
            '"alternatives": ["alt1", "alt2"], '
            '"confidence": "high|medium|low", '
            '"reasoning": "brief explanation only if confidence is NOT high"}'
        )
        instruction = (
            "You are a DE→EN patent terminology expert. "
            "For each German noun, provide: the best English translation for patent context, "
            "1-2 alternatives if they exist, confidence level, and brief reasoning ONLY if "
            "confidence is medium or low. Return a JSON array."
        )
    else:  # adjective
        schema = (
            '{"de": "German", "en": "best translation", '
            '"alternatives": ["alt1", "alt2"], '
            '"avoid": "common mistranslation to avoid (only if confident)", '
            '"confidence": "high|medium|low", '
            '"reasoning": "brief explanation of why literal translation is wrong"}'
        )
        instruction = (
            "You are a DE→EN patent terminology expert. "
            "For each German technical adjective, provide: the best English translation, "
            "1-2 alternatives, an 'avoid' mistranslation ONLY if the literal translation is "
            "definitely wrong in patent context, confidence level, and reasoning. Return a JSON array."
        )

    for batch_start in range(0, len(terms), 40):
        batch = terms[batch_start:batch_start + 40]
        terms_json = json.dumps(batch, ensure_ascii=False)

        prompt = f"{instruction}\n\nSchema per item: {schema}\n\nTerms: {terms_json}"

        try:
            msg = client.messages.create(model=model, max_tokens=8192,
                                         messages=[{"role": "user", "content": prompt}])
            raw = msg.content[0].text.strip()
            code_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', raw, re.DOTALL)
            if code_match:
                raw = code_match.group(1).strip()
            bs = raw.find('[')
            be = raw.rfind(']')
            if bs >= 0 and be > bs:
                parsed = json.loads(raw[bs:be + 1])
                for item in parsed:
                    all_proposals[item["de"]] = item
        except Exception as e:
            pass

    return all_proposals


# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════

tab_terms, tab_translate, tab_review, tab_export = st.tabs([
    "Terminology", "Translate & QE", "Review", "Export"
])


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1: TERMINOLOGY
# ══════════════════════════════════════════════════════════════════════════════

with tab_terms:

    # Pre-compute sentence count for the frequency slider
    if not st.session_state.get("doc_sentence_count"):
        _raw = extract_text_from_docx_terms(docx_path)
        _sents = extract_sentences(_raw)
        st.session_state["doc_sentence_count"] = len(_sents)

    n_sents = st.session_state["doc_sentence_count"]
    auto_default = max(2, min(8, round(n_sents / 20)))
    slider_max = max(auto_default + 3, 10)

    st.markdown(f"**Frequency thresholds** — {n_sents} sentences detected")

    with st.popover("ℹ️ What is the terminology scan?"):
        st.markdown(
            "**Terminology Scan** extracts German technical terms (nouns, adjectives, verbs) "
            "from your patent and proposes English translations for each one. These translations "
            "are then passed to the translator as a glossary so terminology stays consistent "
            "across the entire document.\n\n"
            "**Frequency thresholds** control the minimum number of times a term must appear "
            "in the document to be included. A higher threshold filters out rare or one-off words "
            "and keeps only the core terminology that repeats across claims and description. "
            "A lower threshold captures more terms but may include noise (proper nouns, "
            "one-off compounds, etc.).\n\n"
            "The default is auto-calculated based on document length."
        )

    freq_c1, freq_c2, freq_c3 = st.columns(3)
    with freq_c1:
        min_noun_freq = st.slider(
            "Nouns", min_value=1,
            max_value=slider_max, value=auto_default,
            help=f"Recommended: {auto_default} for {n_sents} sentences",
        )
    with freq_c2:
        min_adj_freq = st.slider(
            "Adjectives", min_value=1, max_value=10, value=2,
            help="Technical adjectives with this many occurrences or more",
        )
    with freq_c3:
        min_verb_freq = st.slider(
            "Verbs", min_value=1, max_value=10, value=2,
            help="Patent verbs with this many occurrences or more",
        )

    # Show scan button only if scan hasn't been done yet
    scan_done = st.session_state.get("noun_counts") is not None
    if scan_done:
        st.success("Terminology scan complete. Review terms below, then proceed to Translate & QE.")
        rescan_col, _ = st.columns([1, 3])
        with rescan_col:
            rescan = st.button("Re-scan", type="secondary")
    else:
        rescan = False

    if (not scan_done and st.button("Scan & Translate Terminology", type="primary", use_container_width=True)) or rescan:
      try:
        import time as _time
        _scan_start = _time.time()
        progress = st.progress(0)
        eta_text = st.empty()

        def _scan_eta(pct):
            elapsed = _time.time() - _scan_start
            if pct > 0.05:
                remaining = elapsed / pct * (1 - pct)
                mins, secs = divmod(int(remaining), 60)
                eta_text.caption(f"~{mins}m {secs}s remaining" if mins else f"~{secs}s remaining")
            progress.progress(min(pct, 1.0))

        _scan_eta(0.02)

        raw_text = extract_text_from_docx_terms(docx_path)
        sentences = extract_sentences(raw_text)
        st.session_state.raw_text = raw_text
        st.session_state.sentences = sentences

        # Diagnostic: detect if text looks like English instead of German
        if sentences:
            sample = " ".join(sentences[:5])
            mid_caps = sum(1 for w in sample.split() if w[0].isupper() and w not in ("FIG.", "Fig."))
            total_words = len(sample.split())
            cap_ratio = mid_caps / total_words if total_words else 0
            if cap_ratio < 0.15:
                st.warning(
                    f"⚠️ Low capitalization ratio ({cap_ratio:.0%}) — this text may be English, "
                    f"not German. The term extractor requires German source text. "
                    f"First 100 chars: `{sample[:100]}…`"
                )

        _scan_eta(0.15)

        if HAS_SPACY:
            try:
                noun_counts, lemma_map = extract_nouns_spacy(sentences)
            except Exception:
                noun_counts = extract_nouns_heuristic(sentences)
                lemma_map = {}
        else:
            noun_counts = extract_nouns_heuristic(sentences)
            lemma_map = {}
        noun_counts = {k: v for k, v in noun_counts.items() if v >= min_noun_freq}

        _scan_eta(0.25)
        adj_freq, adj_variants = extract_technical_adjectives(sentences, min_freq=min_adj_freq)
        verb_info = extract_patent_verbs(sentences)
        # Filter verbs by frequency threshold
        verb_info = {k: v for k, v in verb_info.items() if v.get("frequency", 0) >= min_verb_freq}

        st.session_state.noun_counts = noun_counts
        st.session_state.adj_counts = adj_freq
        st.session_state.adj_variants = adj_variants
        st.session_state.verb_info = verb_info
        st.session_state.lemma_map = lemma_map

        # API translations: nouns
        untranslated_nouns = [n for n in noun_counts if n not in pre_glossary]
        if untranslated_nouns and api_key:
            _scan_eta(0.35)
            noun_proposals = translate_terms_batch(untranslated_nouns, "noun")
            st.session_state.noun_proposals = noun_proposals
            for de, info in noun_proposals.items():
                if de not in st.session_state.glossary:
                    st.session_state.glossary[de] = info.get("en", "")

        # API translations: adjectives
        untranslated_adjs = list(adj_freq.keys())
        if untranslated_adjs and api_key:
            _scan_eta(0.65)
            adj_proposals = translate_terms_batch(untranslated_adjs, "adjective")
            st.session_state.adj_proposals = adj_proposals

        progress.progress(1.0)
        eta_text.empty()
        st.success(f"{len(noun_counts)} nouns · {len(adj_freq)} adjectives · {len(verb_info)} verbs")
      except Exception as e:
        st.error(f"Scan failed: {e}")
        import traceback
        st.code(traceback.format_exc())

    # --- Results in sub-tabs ---
    if st.session_state.noun_counts is not None:
        term_tab1, term_tab2, term_tab3 = st.tabs(["Nouns", "Technical Adjectives", "Patent Verbs"])

        # ---- NOUNS ----
        with term_tab1:
            nouns = st.session_state.noun_counts
            proposals = st.session_state.noun_proposals or {}

            st.caption(f"{len(nouns)} terms — select preferred translation or type your own")

            # Header row
            hc1, hc2, hc3, hc4 = st.columns([2.5, 3, 2, 2])
            hc1.markdown("**German term**")
            hc2.markdown("**English translation**")
            hc3.markdown("**Confidence**")
            hc4.markdown("**Reasoning**")
            st.divider()

            for noun, count in sorted(nouns.items(), key=lambda x: -x[1]):
                locked = noun in pre_glossary
                info = proposals.get(noun, {})
                current_en = st.session_state.glossary.get(noun, info.get("en", ""))
                confidence = info.get("confidence", "")
                alternatives = info.get("alternatives", [])
                reasoning = info.get("reasoning", "")

                col_de, col_en, col_conf, col_reason = st.columns([2.5, 3, 2, 2])

                with col_de:
                    st.markdown(f"{'🔒 ' if locked else ''}**{noun}** ({count}x)")

                with col_en:
                    if locked:
                        st.text(current_en)
                    else:
                        # Build options: current + alternatives + custom
                        options = [current_en] if current_en else []
                        for alt in alternatives:
                            if alt and alt not in options:
                                options.append(alt)
                        if not options:
                            options = [""]
                        options.append("✏️ Custom...")

                        selected = st.selectbox(
                            f"sel_{noun}", options=options,
                            label_visibility="collapsed", key=f"noun_sel_{noun}",
                        )

                        if selected == "✏️ Custom...":
                            custom = st.text_input(
                                f"custom_{noun}", value="", key=f"noun_custom_{noun}",
                                label_visibility="collapsed", placeholder="Type translation..."
                            )
                            if custom:
                                st.session_state.glossary[noun] = custom
                        elif selected:
                            st.session_state.glossary[noun] = selected

                with col_conf:
                    if confidence:
                        css_class = f"conf-{confidence}"
                        st.markdown(f"<span class='{css_class}'>● {confidence}</span>",
                                    unsafe_allow_html=True)

                with col_reason:
                    if reasoning and confidence in ("medium", "low"):
                        st.markdown(f"<span class='reasoning-text'>{reasoning}</span>",
                                    unsafe_allow_html=True)

            # --- Manual glossary entry ---
            st.divider()
            st.markdown("**Add custom term**")
            add_col1, add_col2, add_col3 = st.columns([3, 3, 1])
            with add_col1:
                new_de = st.text_input("German term", key="add_de",
                                       label_visibility="collapsed",
                                       placeholder="German term...")
            with add_col2:
                new_en = st.text_input("English translation", key="add_en",
                                       label_visibility="collapsed",
                                       placeholder="English translation...")
            with add_col3:
                if st.button("➕", key="add_glossary_btn", use_container_width=True):
                    if new_de and new_en:
                        st.session_state.glossary[new_de.strip()] = new_en.strip()
                        # Also add to noun_counts so it shows in the list
                        if st.session_state.noun_counts is not None:
                            if new_de.strip() not in st.session_state.noun_counts:
                                st.session_state.noun_counts[new_de.strip()] = 0
                        st.rerun()

        # ---- ADJECTIVES ----
        with term_tab2:
            adjs = st.session_state.adj_counts or {}
            adj_proposals = st.session_state.adj_proposals or {}

            st.caption(f"{len(adjs)} technical adjectives — review translations and avoid notes")

            hc1, hc2, hc3, hc4 = st.columns([2.5, 3, 2, 2])
            hc1.markdown("**German adjective**")
            hc2.markdown("**English translation**")
            hc3.markdown("**Avoid**")
            hc4.markdown("**Reasoning**")
            st.divider()

            for adj, count in sorted(adjs.items(), key=lambda x: -x[1]):
                info = adj_proposals.get(adj, {})
                en_val = info.get("en", "")
                alternatives = info.get("alternatives", [])
                avoid_val = info.get("avoid", "")
                reasoning = info.get("reasoning", "")
                confidence = info.get("confidence", "")

                col_de, col_en, col_avoid, col_reason = st.columns([2.5, 3, 2, 2])

                with col_de:
                    conf_class = f"conf-{confidence}" if confidence else ""
                    conf_dot = f"<span class='{conf_class}'>● </span>" if confidence else ""
                    st.markdown(f"{conf_dot}**{adj}** ({count}x)", unsafe_allow_html=True)

                with col_en:
                    options = [en_val] if en_val else []
                    for alt in alternatives:
                        if alt and alt not in options:
                            options.append(alt)
                    if not options:
                        options = [""]
                    options.append("✏️ Custom...")

                    selected = st.selectbox(
                        f"adj_sel_{adj}", options=options,
                        label_visibility="collapsed", key=f"adj_sel_{adj}",
                    )
                    if selected == "✏️ Custom...":
                        custom = st.text_input(
                            f"adj_custom_{adj}", value="", key=f"adj_custom_{adj}",
                            label_visibility="collapsed", placeholder="Type translation..."
                        )
                        if custom:
                            st.session_state.glossary[adj] = custom
                    elif selected:
                        st.session_state.glossary[adj] = selected

                with col_avoid:
                    if avoid_val:
                        st.markdown(f"~~{avoid_val}~~")
                    edit_avoid = st.text_input(
                        f"avoid_{adj}", value=avoid_val or "",
                        key=f"adj_avoid_{adj}", label_visibility="collapsed",
                        placeholder="Add avoid term..."
                    )
                    if edit_avoid != (avoid_val or "") and adj in adj_proposals:
                        adj_proposals[adj]["avoid"] = edit_avoid
                        st.session_state.adj_proposals[adj]["avoid"] = edit_avoid

                with col_reason:
                    if reasoning:
                        st.markdown(f"<span class='reasoning-text'>{reasoning}</span>",
                                    unsafe_allow_html=True)

        # ---- VERBS ----
        with term_tab3:
            verbs = st.session_state.verb_info or {}

            st.caption(f"{len(verbs)} patent verbs — add/remove translations individually")

            # Initialize verb_translations in session state from existing data
            if "verb_translations" not in st.session_state:
                st.session_state.verb_translations = {}
                for v, vinfo in verbs.items():
                    raw = vinfo.get("translations", "")
                    # Parse existing slash-separated into list
                    items = [t.strip() for t in raw.replace("/", ",").split(",") if t.strip()]
                    st.session_state.verb_translations[v] = items

            hc1, hc2, hc3 = st.columns([2.5, 4.5, 2])
            hc1.markdown("**German verb**")
            hc2.markdown("**Translations**")
            hc3.markdown("**Forms found**")
            st.divider()

            for verb, info in sorted(verbs.items(), key=lambda x: -x[1]["frequency"]):
                col_v, col_t, col_f = st.columns([2.5, 4.5, 2])

                # Ensure this verb has a list
                if verb not in st.session_state.verb_translations:
                    raw = info.get("translations", "")
                    st.session_state.verb_translations[verb] = [
                        t.strip() for t in raw.replace("/", ",").split(",") if t.strip()
                    ]
                current_tags = st.session_state.verb_translations[verb]

                with col_v:
                    st.markdown(f"**{verb}** ({info['frequency']}x)")

                with col_t:
                    # Show existing translations as removable tags
                    if current_tags:
                        tag_cols = st.columns(min(len(current_tags) + 1, 6))
                        for ti, tag in enumerate(current_tags):
                            with tag_cols[ti % (len(tag_cols) - 1)] if len(tag_cols) > 1 else tag_cols[0]:
                                if st.button(f"✕ {tag}", key=f"rm_verb_{verb}_{ti}",
                                             type="secondary"):
                                    st.session_state.verb_translations[verb].pop(ti)
                                    # Sync back
                                    st.session_state.verb_info[verb]["translations"] = \
                                        ", ".join(st.session_state.verb_translations[verb])
                                    st.rerun()

                    # Add new translation
                    add_c1, add_c2 = st.columns([4, 1])
                    with add_c1:
                        new_trans = st.text_input(
                            f"add_verb_{verb}", value="",
                            label_visibility="collapsed", key=f"verb_add_{verb}",
                            placeholder="Add translation..."
                        )
                    with add_c2:
                        if st.button("➕", key=f"verb_add_btn_{verb}"):
                            if new_trans and new_trans.strip():
                                val = new_trans.strip()
                                if val not in st.session_state.verb_translations[verb]:
                                    st.session_state.verb_translations[verb].append(val)
                                    st.session_state.verb_info[verb]["translations"] = \
                                        ", ".join(st.session_state.verb_translations[verb])
                                    st.rerun()

                with col_f:
                    forms = info.get("forms_found", [])
                    if forms:
                        st.caption(", ".join(forms[:4]))


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2: TRANSLATE & QE
# ══════════════════════════════════════════════════════════════════════════════

with tab_translate:

    if not api_key:
        st.warning("Enter your Anthropic API key in the sidebar.")
        st.stop()
    if not HAS_ANTHROPIC:
        st.error("Install anthropic: pip install anthropic")
        st.stop()

    # Show translate button only if not already done
    translate_done = st.session_state.get("translations") is not None
    if translate_done:
        st.success("Translation & QE complete. Proceed to the Review tab.")
        retrans_col, _ = st.columns([1, 3])
        with retrans_col:
            retranslate = st.button("Re-translate", type="secondary")
    else:
        retranslate = False

    if (not translate_done and st.button("Translate & Evaluate", type="primary", use_container_width=True)) or retranslate:
        # Demo mode: check remaining patents
        if DEMO_MODE:
            recheck = validate_code(st.session_state.demo_code)
            if not recheck["valid"]:
                st.error(recheck["message"])
                st.stop()
            record_patent_use(st.session_state.demo_code, st.session_state.get("docx_name", "patent.docx"))

        import time as _time
        _pipe_start = _time.time()
        progress = st.progress(0)
        eta_text = st.empty()

        def _pipe_eta(pct):
            elapsed = _time.time() - _pipe_start
            if pct > 0.03:
                remaining = elapsed / pct * (1 - pct)
                mins, secs = divmod(int(remaining), 60)
                eta_text.caption(f"~{mins}m {secs}s remaining" if mins else f"~{secs}s remaining")
            progress.progress(min(pct, 1.0))

        _pipe_eta(0.02)

        # Step 1: Structural analysis (silent)
        nlp = load_nlp()
        if nlp:
            structural = analyze_document(docx_path, nlp)
            st.session_state.structural_analysis = structural
        else:
            st.session_state.structural_analysis = None

        _pipe_eta(0.08)

        # Step 2: Translate
        raw_text = extract_text_translate(docx_path)
        sentences = split_sentences_translate(raw_text)
        glossary = st.session_state.glossary if st.session_state.glossary else None

        def trans_progress(current, total):
            pct = 0.1 + (current / total) * 0.5
            _pipe_eta(pct)

        results = translate_document(
            sentences=sentences, api_key=api_key, model=model,
            glossary=glossary, structural_analysis=st.session_state.structural_analysis,
            batch_size=batch_size_trans, progress_callback=trans_progress,
            domain=selected_domain,
        )

        st.session_state.translations = {
            "metadata": {
                "source_file": st.session_state.get("docx_name", "patent.docx"), "model": model,
                "n_sentences": len(sentences),
                "n_structural_hints": sum(1 for r in results if r["had_structural_hint"]),
                "glossary_terms": len(st.session_state.glossary),
            },
            "translations": results,
        }

        # Step 3: QE
        _pipe_eta(0.62)

        few_shot = None
        if training_pairs_path and os.path.exists(training_pairs_path) and n_few_shot > 0:
            all_pairs = load_training_pairs(training_pairs_path)
            few_shot = select_few_shot_examples(all_pairs, n_few_shot)

        def qe_progress(current, total):
            pct = 0.65 + (current / total) * 0.3
            _pipe_eta(pct)

        qe_results = evaluate_translations(
            translations=results, api_key=api_key, model=model,
            few_shot_examples=few_shot, batch_size=batch_size_qe,
            progress_callback=qe_progress, domain=selected_domain,
        )

        triage = compute_triage(qe_results)
        st.session_state.qe_results = qe_results
        st.session_state.triage = triage

        # Check for QE failures and warn user
        n_failed = sum(1 for r in qe_results if "[QE FAILED" in r.get("explanation", ""))
        if n_failed > 0:
            st.warning(
                f"⚠️ QE could not evaluate {n_failed}/{len(qe_results)} segments. "
                f"These are marked as 'review' — check the Streamlit terminal for error details. "
                f"Common causes: API rate limits, model unavailability, or very long segments."
            )

        progress.progress(1.0)
        eta_text.empty()

    # --- Results ---
    if st.session_state.triage:
        triage = st.session_state.triage
        s = triage["summary"]

        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"""<div class="triage-card triage-green">
                <div style="font-size:2.2rem">{s['green_count']}</div>
                <div>Publishable ({s['green_pct']:.0f}%)</div>
            </div>""", unsafe_allow_html=True)
        with col2:
            st.markdown(f"""<div class="triage-card triage-orange">
                <div style="font-size:2.2rem">{s['orange_count']}</div>
                <div>Review ({s['orange_pct']:.0f}%)</div>
            </div>""", unsafe_allow_html=True)
        with col3:
            st.markdown(f"""<div class="triage-card triage-red">
                <div style="font-size:2.2rem">{s['red_count']}</div>
                <div>Edit needed ({s['red_pct']:.0f}%)</div>
            </div>""", unsafe_allow_html=True)

        st.markdown("")

        error_breakdown = s.get("error_breakdown", {})
        if error_breakdown:
            cols = st.columns(min(len(error_breakdown), 6))
            for i, (etype, count) in enumerate(sorted(error_breakdown.items(), key=lambda x: -x[1])):
                with cols[i % len(cols)]:
                    st.metric(etype.title(), count)

        if s.get("high_risk_total", 0) > 0:
            hr_pct = round(s["high_risk_green"] / s["high_risk_total"] * 100) if s["high_risk_total"] else 0
            st.caption(f"Structural analysis: {s['high_risk_total']} high-risk sentences → "
                       f"{s['high_risk_green']} publishable ({hr_pct}%)")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3: REVIEW
# ══════════════════════════════════════════════════════════════════════════════

with tab_review:

    if st.session_state.translations is None or st.session_state.qe_results is None:
        st.info("Run Translate & QE first.")
        st.stop()

    translations = st.session_state.translations["translations"]
    qe_by_idx = {r["index"]: r for r in st.session_state.qe_results}
    green_set = set(st.session_state.triage["green"]) if st.session_state.triage else set()
    confirmed = st.session_state.confirmed

    # Progress tracker
    total_segs = len(translations)
    n_confirmed = len(confirmed)
    n_green = len(green_set)
    # "Done" = confirmed + unconfirmed greens (publishable without edits)
    n_done = n_confirmed + sum(1 for i in green_set if i not in confirmed)
    pct_confirmed = int(n_confirmed / total_segs * 100) if total_segs else 0
    pct_done = int(n_done / total_segs * 100) if total_segs else 0
    needs_review = total_segs - n_done

    st.progress(pct_confirmed / 100, text=f"**{pct_confirmed}%** confirmed ({n_confirmed}/{total_segs})  ·  {needs_review} segments need review")

    # Filter controls
    col_filter, _ = st.columns([4, 1])
    with col_filter:
        view_mode = st.radio(
            "Show", ["All segments", "Needs review", "Red only", "Confirmed"],
            horizontal=True, label_visibility="collapsed",
        )

    st.divider()

    @st.fragment
    def _review_segments():
        """Fragment-wrapped segment list — reruns only this section on button clicks."""
        _translations = st.session_state.translations["translations"]
        _qe_by_idx = {r["index"]: r for r in st.session_state.qe_results}
        _green_set = set(st.session_state.triage["green"]) if st.session_state.triage else set()
        _confirmed = st.session_state.confirmed

        for trans in _translations:
            idx = trans["index"]
            qe = _qe_by_idx.get(idx, {})
            rating = qe.get("rating", "unknown")
            is_green = idx in _green_set
            is_confirmed = idx in _confirmed
            suggestion = qe.get("suggestion", "")
            explanation = qe.get("explanation", "")
            error_cat = qe.get("error_category", "")

            # Filter logic
            if view_mode == "Needs review" and (is_green or is_confirmed):
                continue
            elif view_mode == "Red only" and rating not in ("major", "critical"):
                continue
            elif view_mode == "Confirmed" and not is_confirmed:
                continue

            badge = {"good": "🟢", "minor": "🟡", "major": "🔴", "critical": "⛔"}.get(rating, "⚪")

            # --- Segment header ---
            header_col, action_col = st.columns([5, 2])
            with header_col:
                st.markdown(f"{badge} **Segment {idx}** — _{rating}_")
            with action_col:
                btn_cols = st.columns(2)
                if is_confirmed:
                    with btn_cols[0]:
                        st.markdown("✅ Locked")
                    with btn_cols[1]:
                        def _unlock(i=idx):
                            del st.session_state.confirmed[i]
                        st.button("Unlock", key=f"unlock_{idx}", type="secondary", on_click=_unlock)
                else:
                    with btn_cols[0]:
                        def _confirm(i=idx, t=trans):
                            edited = st.session_state.get(f"edit_{i}", t.get("translation", ""))
                            st.session_state.confirmed[i] = edited
                        st.button("Confirm", key=f"confirm_{idx}", type="secondary", on_click=_confirm)
                    with btn_cols[1]:
                        # Apply QE suggestion — puts text in editor, does NOT lock
                        if suggestion and rating != "good":
                            already_applied = st.session_state.get(f"applied_{idx}", False)
                            if already_applied:
                                st.markdown("✅ Fix applied")
                            else:
                                def _apply_fix(i=idx, sug=suggestion):
                                    clean_sug = re.sub(r'\*\*(.+?)\*\*', r'\1', sug)
                                    st.session_state[f"edit_{i}"] = clean_sug
                                    st.session_state[f"applied_{i}"] = True
                                st.button("Apply fix", key=f"apply_{idx}", type="primary", on_click=_apply_fix)

            # --- Side by side ---
            col_de, col_en = st.columns(2)

            with col_de:
                st.caption("German source")
                st.markdown(f"<div class='source-box'>{trans.get('source', '')}</div>",
                            unsafe_allow_html=True)

            with col_en:
                st.caption("English translation")
                if is_confirmed:
                    st.markdown(f"<div class='locked-box'>{_confirmed[idx]}</div>",
                                unsafe_allow_html=True)
                else:
                    # Only set default value if key not already in session state
                    # (Apply fix / Add term sets it via session state API)
                    if f"edit_{idx}" not in st.session_state:
                        st.session_state[f"edit_{idx}"] = trans.get("translation", "")
                    st.text_area(
                        f"edit_{idx}",
                        label_visibility="collapsed", key=f"edit_{idx}",
                        height=100,
                    )

            # --- QE reasoning ---
            if rating != "good" and explanation:
                qe_class = {"major": "qe-major", "critical": "qe-critical"}.get(rating, "qe-minor")
                suggestion_html = ""
                if suggestion:
                    # Convert **bold** markers to <strong> for HTML rendering
                    rendered = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', suggestion)
                    suggestion_html = (
                        f"<div class='qe-suggestion'>"
                        f"<strong>Suggested:</strong> {rendered}"
                        f"</div>"
                    )
                st.markdown(
                    f"<div class='qe-note {qe_class}'>"
                    f"<strong>{error_cat}:</strong> {explanation}"
                    f"{suggestion_html}"
                    f"</div>", unsafe_allow_html=True
                )
            elif rating == "good":
                st.markdown(
                    f"<div class='qe-note qe-good'>No issues detected</div>",
                    unsafe_allow_html=True
                )

            st.markdown("---")

    _review_segments()


# ══════════════════════════════════════════════════════════════════════════════
# TAB 4: EXPORT
# ══════════════════════════════════════════════════════════════════════════════

with tab_export:

    if st.session_state.translations is None or st.session_state.qe_results is None:
        st.info("Run Translate & QE first.")
        st.stop()

    st.markdown("### QE Review Document")
    st.caption("Color-coded .docx with triage summary, source text, and QE annotations")

    col_opt1, col_opt2 = st.columns(2)
    with col_opt1:
        include_source = st.checkbox("Include German source", value=True)
    with col_opt2:
        include_annotations = st.checkbox("Include QE notes", value=True)

    if st.button("Generate QE Document", type="primary"):
        with st.spinner("Building document..."):
            qe_data = {
                "qe_results": st.session_state.qe_results,
                "triage": st.session_state.triage,
            }
            output_path = os.path.join(tempfile.gettempdir(), "glosswerk_qe_review.docx")
            assemble_document(
                translations_data=st.session_state.translations,
                qe_data=qe_data, output_path=output_path,
                include_source=include_source,
                include_annotations=include_annotations,
            )
            with open(output_path, "rb") as f:
                st.download_button(
                    "Download QE Review Document", data=f.read(),
                    file_name=f"glosswerk_qe_{st.session_state.get('docx_name', 'patent.docx')}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

    st.divider()

    st.markdown("### Translation Only")
    st.caption("Clean English document — uses confirmed edits where available")

    if st.button("Generate Translation Document"):
        with st.spinner("Building document..."):
            from docx import Document as DocxDoc

            doc = DocxDoc()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(10)

            for trans in st.session_state.translations["translations"]:
                idx = trans["index"]
                text = st.session_state.confirmed.get(idx, trans.get("translation", ""))
                if text:
                    doc.add_paragraph(text)

            # Demo watermark
            if DEMO_MODE:
                from docx.shared import RGBColor
                wp = doc.add_paragraph()
                wp.alignment = 1  # center
                run = wp.add_run(WATERMARK_TEXT)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.italic = True

            out_path = os.path.join(tempfile.gettempdir(), "glosswerk_translation.docx")
            doc.save(out_path)

            with open(out_path, "rb") as f:
                st.download_button(
                    "Download Translation", data=f.read(),
                    file_name=f"translation_{st.session_state.get('docx_name', 'patent.docx')}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

    st.divider()

    st.markdown("### Data Exports")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Translations (JSON)",
            data=json.dumps(st.session_state.translations, ensure_ascii=False, indent=2),
            file_name="translations.json", mime="application/json",
        )
    with col2:
        qe_export = {"qe_results": st.session_state.qe_results, "triage": st.session_state.triage}
        st.download_button(
            "QE Results (JSON)",
            data=json.dumps(qe_export, ensure_ascii=False, indent=2),
            file_name="qe_results.json", mime="application/json",
        )

    # Merge scan glossary + review-time additions for export
    combined_glossary = dict(st.session_state.glossary) if st.session_state.glossary else {}
    review_additions = st.session_state.get("user_glossary_additions", {})
    if review_additions:
        # Review additions use English find→replace, store as-is for reference
        combined_glossary.update({f"[review] {k}": v for k, v in review_additions.items()})

    if combined_glossary:
        st.divider()
        st.markdown("### Glossary Export")
        n_scan = len(st.session_state.glossary) if st.session_state.glossary else 0
        n_review = len(review_additions)
        st.caption(f"{n_scan} terms from scan · {n_review} terms added during review")

        glossary_tsv_lines = []
        # Scan terms (DE → EN)
        if st.session_state.glossary:
            glossary_tsv_lines.append("# Terms from terminology scan (DE → EN)")
            for de, en in sorted(st.session_state.glossary.items()):
                glossary_tsv_lines.append(f"{de}\t{en}")
        # Review additions (EN find → EN replace)
        if review_additions:
            glossary_tsv_lines.append("")
            glossary_tsv_lines.append("# Terms added during review (find → replace)")
            for find_t, replace_t in sorted(review_additions.items()):
                glossary_tsv_lines.append(f"{find_t}\t{replace_t}")

        glossary_tsv = "\n".join(glossary_tsv_lines)
        st.download_button(
            "Glossary (TSV)", data=glossary_tsv,
            file_name="glossary.tsv", mime="text/tab-separated-values",
        )
