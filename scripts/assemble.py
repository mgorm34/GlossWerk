"""
GlossWerk Document Assembly & Triage Module

Takes translation + QE results and produces:
1. Color-coded .docx with green/orange/red segments
2. Triage summary page at the top
3. Side-by-side German source + English translation
4. QE annotations (error type, explanation) for non-green segments

Usage:
    python assemble.py --translations translations.json \
        --qe qe_results.json --output patent_reviewed.docx

Requires: python-docx
Install:  pip install python-docx
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Color definitions
# ---------------------------------------------------------------------------

COLOR_GREEN = RGBColor(0x22, 0x8B, 0x22)     # ForestGreen
COLOR_ORANGE = RGBColor(0xFF, 0x8C, 0x00)    # DarkOrange
COLOR_RED = RGBColor(0xCC, 0x00, 0x00)       # Dark red
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)      # Gray for source text
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)

BG_GREEN = "D5F5D5"    # Light green background
BG_ORANGE = "FFF3CD"   # Light orange/yellow background
BG_RED = "F8D7DA"      # Light red background


def get_triage_color(rating):
    """Get the color for a QE rating."""
    if rating == "good":
        return COLOR_GREEN
    elif rating == "minor":
        return COLOR_ORANGE
    else:  # major, critical
        return COLOR_RED


def get_triage_bg(rating):
    """Get background shading hex for a rating."""
    if rating == "good":
        return BG_GREEN
    elif rating == "minor":
        return BG_ORANGE
    else:
        return BG_RED


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.find(qn('w:shd'))
    if shading_elem is None:
        from lxml import etree
        shading_elem = etree.SubElement(shading, qn('w:shd'))
    shading_elem.set(qn('w:fill'), color_hex)
    shading_elem.set(qn('w:val'), 'clear')


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def assemble_document(translations_data, qe_data, output_path,
                      include_source=True, include_annotations=True):
    """
    Assemble the final color-coded .docx document.

    Args:
        translations_data: dict from translate.py output
        qe_data: dict from quality_estimate.py output
        output_path: path for output .docx
        include_source: include German source column
        include_annotations: include QE annotation column
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # --- Title page / header ---
    metadata = translations_data.get("metadata", {})
    triage = qe_data.get("triage", {})
    summary = triage.get("summary", {})

    title = doc.add_heading('GlossWerk Translation Review', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Source file info
    source_file = metadata.get("source_file", "Unknown")
    model = metadata.get("model", "Unknown")
    n_sentences = metadata.get("n_sentences", 0)
    n_hints = metadata.get("n_structural_hints", 0)
    glossary_n = metadata.get("glossary_terms", 0)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        f"Source: {source_file}\n"
        f"Model: {model} | Segments: {n_sentences} | "
        f"Structural hints: {n_hints} | Glossary terms: {glossary_n}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_GRAY

    # --- Triage summary box ---
    doc.add_heading('Triage Summary', level=2)

    # Summary table
    triage_table = doc.add_table(rows=4, cols=4)
    triage_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Category', 'Count', 'Percentage', 'Action']
    for i, h in enumerate(headers):
        cell = triage_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Green row
    row = triage_table.rows[1]
    _set_summary_row(row, 'GREEN (Publishable)',
                     summary.get('green_count', 0),
                     summary.get('green_pct', 0),
                     'No review needed', BG_GREEN)

    # Orange row
    row = triage_table.rows[2]
    _set_summary_row(row, 'ORANGE (Quick Review)',
                     summary.get('orange_count', 0),
                     summary.get('orange_pct', 0),
                     'Minor fixes possible', BG_ORANGE)

    # Red row
    row = triage_table.rows[3]
    _set_summary_row(row, 'RED (Full Edit)',
                     summary.get('red_count', 0),
                     summary.get('red_pct', 0),
                     'Significant editing needed', BG_RED)

    doc.add_paragraph()  # Spacing

    # Error breakdown if available
    error_breakdown = summary.get("error_breakdown", {})
    if error_breakdown:
        doc.add_heading('Error Types (non-green segments)', level=3)
        for etype, count in sorted(error_breakdown.items(), key=lambda x: -x[1]):
            doc.add_paragraph(f"{etype}: {count}", style='List Bullet')

    # Structural risk cross-reference
    hr_total = summary.get("high_risk_total", 0)
    hr_green = summary.get("high_risk_green", 0)
    if hr_total > 0:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            f"Structural analysis: {hr_total} high-risk sentences identified. "
            f"Of those, {hr_green} were translated well enough to be GREEN "
            f"({round(hr_green/hr_total*100, 1) if hr_total else 0}% — "
            f"structural hints helped guide translation)."
        )
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_page_break()

    # --- Translation segments ---
    doc.add_heading('Translation Segments', level=2)

    legend = doc.add_paragraph()
    legend_run = legend.add_run('Legend: ')
    legend_run.font.bold = True
    legend_run.font.size = Pt(9)

    for label, color in [('GREEN = publishable', COLOR_GREEN),
                         (' | ORANGE = review', COLOR_ORANGE),
                         (' | RED = edit needed', COLOR_RED)]:
        run = legend.add_run(label)
        run.font.color.rgb = color
        run.font.size = Pt(9)
        run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Build QE lookup
    qe_results = qe_data.get("qe_results", [])
    qe_by_idx = {r["index"]: r for r in qe_results}

    translations = translations_data.get("translations", [])

    # Main translation table
    n_cols = 1  # Translation always present
    if include_source:
        n_cols += 1
    if include_annotations:
        n_cols += 1

    # Add segments
    for trans in translations:
        idx = trans["index"]
        source = trans.get("source", "")
        translation = trans.get("translation", "")
        qe = qe_by_idx.get(idx, {})
        rating = qe.get("rating", "minor")
        error_cat = qe.get("error_category", "")
        explanation = qe.get("explanation", "")
        risk_level = trans.get("risk_level", "")
        risk_score = trans.get("risk_score", 0)

        color = get_triage_color(rating)
        bg = get_triage_bg(rating)

        # Segment header with index and rating
        header_para = doc.add_paragraph()
        idx_run = header_para.add_run(f"[{idx}] ")
        idx_run.font.bold = True
        idx_run.font.size = Pt(9)
        idx_run.font.color.rgb = COLOR_GRAY

        rating_run = header_para.add_run(f"{rating.upper()}")
        rating_run.font.bold = True
        rating_run.font.size = Pt(9)
        rating_run.font.color.rgb = color

        if risk_level in ("medium", "high"):
            risk_run = header_para.add_run(f" | risk={risk_level} ({risk_score:.2f})")
            risk_run.font.size = Pt(8)
            risk_run.font.color.rgb = COLOR_GRAY

        # Translation text (color-coded)
        trans_para = doc.add_paragraph()
        trans_run = trans_para.add_run(translation)
        trans_run.font.size = Pt(10)
        trans_run.font.color.rgb = color

        # Source text (gray, smaller)
        if include_source:
            src_para = doc.add_paragraph()
            src_label = src_para.add_run("DE: ")
            src_label.font.bold = True
            src_label.font.size = Pt(8)
            src_label.font.color.rgb = COLOR_GRAY
            src_run = src_para.add_run(source)
            src_run.font.size = Pt(8)
            src_run.font.color.rgb = COLOR_GRAY

        # QE annotation (only for non-green)
        if include_annotations and rating != "good" and explanation:
            ann_para = doc.add_paragraph()
            ann_label = ann_para.add_run(f"QE [{error_cat}]: ")
            ann_label.font.bold = True
            ann_label.font.size = Pt(8)
            ann_label.font.color.rgb = color
            ann_run = ann_para.add_run(explanation)
            ann_run.font.size = Pt(8)
            ann_run.font.color.rgb = color
            ann_run.font.italic = True

        # Thin separator
        sep = doc.add_paragraph()
        sep_run = sep.add_run('─' * 80)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Save
    doc.save(output_path)
    return output_path


def _set_summary_row(row, label, count, pct, action, bg_hex):
    """Set values and formatting for a triage summary table row."""
    cells = row.cells
    cells[0].text = label
    cells[1].text = str(count)
    cells[2].text = f"{pct:.1f}%"
    cells[3].text = action

    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
        set_cell_shading(cell, bg_hex)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    import argparse

    parser = argparse.ArgumentParser(description="Assemble color-coded translation .docx")
    parser.add_argument("--translations", required=True,
                        help="JSON file from translate.py")
    parser.add_argument("--qe", required=True,
                        help="JSON file from quality_estimate.py")
    parser.add_argument("--output", required=True,
                        help="Output .docx file path")
    parser.add_argument("--no-source", action="store_true",
                        help="Omit German source text")
    parser.add_argument("--no-annotations", action="store_true",
                        help="Omit QE annotations")
    args = parser.parse_args()

    # Load data
    with open(args.translations, "r", encoding="utf-8") as f:
        trans_data = json.load(f)

    with open(args.qe, "r", encoding="utf-8") as f:
        qe_data = json.load(f)

    print(f"Loaded {len(trans_data.get('translations', []))} translations")
    print(f"Loaded {len(qe_data.get('qe_results', []))} QE results")

    # Assemble
    output_path = assemble_document(
        translations_data=trans_data,
        qe_data=qe_data,
        output_path=args.output,
        include_source=not args.no_source,
        include_annotations=not args.no_annotations,
    )

    triage = qe_data.get("triage", {}).get("summary", {})
    print(f"\nDocument assembled: {output_path}")
    print(f"  GREEN:  {triage.get('green_count', 0)} ({triage.get('green_pct', 0):.1f}%)")
    print(f"  ORANGE: {triage.get('orange_count', 0)} ({triage.get('orange_pct', 0):.1f}%)")
    print(f"  RED:    {triage.get('red_count', 0)} ({triage.get('red_pct', 0):.1f}%)")


if __name__ == "__main__":
    main()
